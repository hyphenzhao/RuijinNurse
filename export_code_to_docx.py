#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
从项目代码中收集源码，按行编号，按软著要求导出为 Word 文档：
- 若总行数 <= 4000，则全部导出；
- 若总行数 > 4000，则导出前 2000 行 + 后 2000 行；
- Word 文档按每页 50 行插入分页符。

使用方式：
    python export_code_to_docx.py

生成文件：
    神人一号_源代码_前后40页.docx
"""

import os
from pathlib import Path

from docx import Document

# ---------------- 配置区域（可按需要修改） ---------------- #

# 项目根目录（默认当前脚本所在目录）
ROOT_DIR = Path(__file__).resolve().parent

# 需要纳入统计的源代码后缀
INCLUDE_EXTS = {
    ".py", ".cs", ".js", ".html", ".htm", ".css",
    ".cpp", ".cc", ".cxx", ".c", ".h", ".hpp",
    ".json", ".yml", ".yaml", ".xml",
    ".sass", ".scss", ".lua", ".java",
}

# 需要排除的目录（避免 venv、构建目录等）
EXCLUDE_DIRS = {
    ".git", ".idea", ".vscode",
    "venv", ".venv", "env", "ENV",
    "Library", "Build", "Builds", "Logs",
    "__pycache__",
    "node_modules",
}

# 每页行数
LINES_PER_PAGE = 50

# 导出规则相关参数
PAGE_COUNT_FRONT = 40
PAGE_COUNT_BACK = 40
LINES_FRONT = PAGE_COUNT_FRONT * LINES_PER_PAGE  # 2000
LINES_BACK = PAGE_COUNT_BACK * LINES_PER_PAGE    # 2000

# 输出文件名
OUTPUT_DOCX = ROOT_DIR / "神人一号_源代码_前后40页.docx"

# ---------------- 工具函数 ---------------- #

def sanitize_for_xml(s: str) -> str:
    """
    移除 / 替换掉 Word(XML) 不支持的控制字符。
    XML 1.0 允许的字符大致为：
    - 0x9, 0xA, 0xD
    - 0x20 - 0xD7FF
    - 0xE000 - 0xFFFD
    其余全部用空格替换。
    """
    if not isinstance(s, str):
        s = str(s)

    res_chars = []
    for ch in s:
        code = ord(ch)
        if (
            code == 0x9
            or code == 0xA
            or code == 0xD
            or (0x20 <= code <= 0xD7FF)
            or (0xE000 <= code <= 0xFFFD)
        ):
            res_chars.append(ch)
        else:
            # 用空格替代非法字符（也可以用 "" 直接丢弃）
            res_chars.append(" ")
    return "".join(res_chars)


def should_skip_dir(dir_name: str) -> bool:
    """判断是否跳过该目录"""
    return dir_name in EXCLUDE_DIRS


def should_include_file(path: Path) -> bool:
    """判断是否纳入统计的文件类型"""
    return path.suffix.lower() in INCLUDE_EXTS


def collect_code_lines(root: Path):
    """
    遍历目录，收集所有源代码文件的行，返回一个字符串列表。
    每行前面会统一编号（全局递增），方便软著查阅。
    """
    all_lines = []
    global_line_no = 1

    # 按文件路径排序，保证顺序稳定
    file_paths = []
    for dirpath, dirnames, filenames in os.walk(root):
        # 过滤目录
        dirnames[:] = [d for d in dirnames if not should_skip_dir(d)]

        for fname in filenames:
            fpath = Path(dirpath) / fname
            if should_include_file(fpath):
                file_paths.append(fpath)

    file_paths = sorted(file_paths, key=lambda p: str(p))

    for fpath in file_paths:
        rel_path = fpath.relative_to(root)
        header = f"##### FILE: {rel_path} #####"
        # 每个文件前加一个文件头，方便审查看到是哪个文件
        # all_lines.append(f"{global_line_no:06d}  {header}")
        global_line_no += 1

        try:
            with fpath.open("r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    text = line.rstrip("\n\r")
                    # all_lines.append(f"{global_line_no:06d}  {text}")
                    all_lines.append(f"{text}")
                    global_line_no += 1
        except Exception as e:
            # 避免单个文件出问题导致整体失败
            err_msg = f"# [ERROR READING FILE {rel_path}: {e}]"
            all_lines.append(f"{global_line_no:06d}  {err_msg}")
            global_line_no += 1

    return all_lines


def select_lines_for_soft_ip(all_lines):
    """
    根据总行数选择：
    - <= 4000 行：全部导出；
    - > 4000 行：导出前 2000 行 + 后 2000 行。
    """
    total = len(all_lines)
    print(f"总行数: {total}")

    if total <= (LINES_FRONT + LINES_BACK):
        print("总行数不超过 4000 行，将全部导出。")
        return all_lines

    print("总行数超过 4000 行，将导出前 2000 行和后 2000 行。")
    front_part = all_lines[:LINES_FRONT]
    back_part = all_lines[-LINES_BACK:]
    return front_part + back_part


def export_to_docx(lines, output_path: Path, lines_per_page: int = 50):
    """
    将选定的代码行导出到 Word 文档，按每页 lines_per_page 行插入分页符。
    """
    doc = Document()

    # 可选：封面说明
    # doc.add_heading("神人一号 源代码节选（前40页 + 后40页）", level=1)
    # doc.add_paragraph("说明：每页约 50 行，包含项目部分源代码，供软件著作权登记使用。")
    # doc.add_page_break()

    for idx, text in enumerate(lines):
        # 每满 lines_per_page 行插入分页符
        # if idx > 0 and idx % lines_per_page == 0:
        #     doc.add_page_break()

        # 清洗不合法字符，防止 python-docx 报 XML 兼容错误
        safe_text = sanitize_for_xml(text)

        # 每行作为一个段落
        p = doc.add_paragraph(safe_text)


    doc.save(output_path)
    print(f"已生成 Word 文档: {output_path}")


# ---------------- 主执行逻辑 ---------------- #


def main():
    print(f"项目根目录: {ROOT_DIR}")
    print("开始收集代码行...")

    all_lines = collect_code_lines(ROOT_DIR)
    print("代码收集完成。")

    selected_lines = select_lines_for_soft_ip(all_lines)
    print(f"最终导出行数: {len(selected_lines)}")

    export_to_docx(selected_lines, OUTPUT_DOCX, lines_per_page=LINES_PER_PAGE)


if __name__ == "__main__":
    main()
